from datetime import date
from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "CareLedger_AI_Technical_Documentation.docx"

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
for name, size, color in [("Title", 30, "0F766E"), ("Heading 1", 20, "0F766E"), ("Heading 2", 15, "134E4A"), ("Heading 3", 12, "334155")]:
    styles[name].font.name = "Aptos Display"
    styles[name].font.size = Pt(size)
    styles[name].font.color.rgb = RGBColor.from_string(color)

def p(text="", style=None):
    return doc.add_paragraph(text, style)

def bullets(items):
    for item in items:
        p(item, "List Bullet")

def numbered(items):
    for item in items:
        p(item, "List Number")

def code(text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(7)
    para.paragraph_format.left_indent = Inches(0.18)
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), "F1F5F9")
    para._p.get_or_add_pPr().append(shade)
    return para

def heading(text, level=1):
    return doc.add_heading(text, level=level)

def add_toc():
    heading("Table of Contents", 1)
    para = doc.add_paragraph()
    run = para.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "Right-click and choose Update Field if page numbers are not shown."
    separate.append(text)
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])
    bullets([
        "1. Project Overview", "2. Tech Stack", "3. Project Structure",
        "4. Feature-by-Feature Logic", "5. Data Flow Between Features",
        "6. Authentication Flow", "7. The ML Model (Disease Prediction)",
        "8. Known Limitations / Honest Gaps"
    ])

# Title page
for _ in range(5): p()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("CareLedger AI — Technical Documentation")
r.bold = True; r.font.size = Pt(30); r.font.color.rgb = RGBColor(15, 118, 110)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Verified against the repository implementation").italic = True
dt = doc.add_paragraph(); dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
dt.add_run(date.today().strftime("%B %d, %Y"))
doc.add_page_break()
add_toc()
doc.add_page_break()

heading("1. Project Overview")
p("CareLedger AI is a web application that helps a signed-in user move from an initial symptom check to a machine-learning disease prediction, provider guidance, cost planning, appointment booking, follow-up questions through an AI assistant, and a clearly separated emergency screen. It is decision-support software, not a clinical diagnosis or treatment system.")
heading("End-to-end user journey", 2)
numbered([
    "Health Assessment: the user answers nine multiple-choice questions. Browser-side scoring produces a risk percentage, severity, and next-step message. The result can seed the shared care context, but it is not sent to the backend.",
    "Disease Prediction: the user selects at least two symptoms from the backend-provided model feature list. The backend runs a trained RandomForestClassifier, augments its probabilities with a compact disease rule/knowledge base, returns ranked candidates and emergency metadata, and stores the leading result in the Zustand care context.",
    "Smart Care Navigator: it reads the latest disease, specialist, risk, confidence, and source. The backend ranks curated Pakistani hospital and fabricated doctor profiles for the specialty, calculates a consultation-and-labs preview, and optionally uses Gemini to prepare visit guidance.",
    "Cost Intelligence: it pre-fills the predicted disease and selected hospital from Zustand. Deterministic PKR cost bands produce the estimate; the affordability score is formula-based; Gemini is used only for the narrative summary and suggestions when available.",
    "Appointments: a hospital/doctor handoff from Navigator preselects the catalog entry. The user chooses date, slot, visit type, and reason. The backend validates catalog availability and detects conflicts, then stores the booking in memory and returns a confirmation (Gemini-authored when available, deterministic otherwise).",
    "AI Chat: the user asks general health questions. Supabase stores per-user chat sessions/messages, while the FastAPI endpoint sends the current message and conversation history to Gemini. A local keyword detector adds an emergency flag.",
    "Emergency Mode: red flags from Assessment, Disease Prediction, Navigator, or AI Chat can direct the user to a static emergency page with Pakistan Rescue 1122 calling and basic first-aid guidance. Emergency Mode itself makes no API call."
])
p("The pipeline is connected through browser state and navigation handoffs rather than a backend workflow engine. A user may also open features directly; the Navigator explicitly asks for a prior assessment/prediction if no care context exists.")

heading("2. Tech Stack")
heading("Frontend", 2)
bullets([
    "React 18.2 and react-dom render the active web application in frontend/src.",
    "Vite 5 provides the development server and production build; @vitejs/plugin-react handles React transforms.",
    "Tailwind CSS 3.4, PostCSS, and Autoprefixer provide the design system and utility styling.",
    "react-router-dom 6.28 defines public, guest-only, and protected routes and carries short-lived navigation state between features.",
    "Zustand 4.5 with persist middleware stores the care pipeline in sessionStorage.",
    "@supabase/supabase-js 2.111 handles authentication and direct browser access to chat history tables.",
    "@heroicons/react provides interface icons. Native fetch performs backend calls; there is no Axios dependency."
])
heading("Backend", 2)
bullets([
    "FastAPI exposes synchronous JSON route handlers; Pydantic validates request/response models; Uvicorn serves the ASGI app.",
    "pandas builds a one-row feature DataFrame and reads training/testing CSV files; NumPy constructs the binary symptom vector.",
    "scikit-learn supplies RandomForestClassifier and accuracy_score; joblib loads and saves the model and label/feature artifacts.",
    "google-genai (not google-generativeai) integrates Gemini for AI Chat, Navigator visit guidance, affordability narrative, and appointment confirmation.",
    "python-dotenv loads the repository-level .env before routers are imported."
])
heading("Database and authentication", 2)
p("Supabase is used for authentication and chat persistence. Auth sessions are stored/refreshed by the Supabase browser client. The SQL migration creates chat_sessions and chat_messages with row-level security tied to auth.uid(). Supabase is not used by the FastAPI endpoints, is not a disease-record database, and does not store appointments. Profile images are browser localStorage data.")
heading("Gemini integration", 2)
bullets([
    "AI Chat: a safety-focused health-assistant system instruction requests concise plain-language answers, prohibits definitive diagnosis/prescribing, and directs emergencies to urgent care. Conversation roles are normalized before submission.",
    "Smart Care Navigator: asks for JSON containing three short what-to-expect items and four questions to ask before a visit; deterministic fallback text is available.",
    "Affordability: receives the real computed costs, income, savings, insurance, and score and is asked for a short empathetic summary plus exactly three number-specific suggestions; deterministic fallbacks remain authoritative if Gemini fails.",
    "Appointments: asks for a warm confirmation under three sentences that names the selected doctor, hospital, date, time, visit type, fee, and optionally the visit reason."
])
p("All four integrations use GEMINI_API_KEY and default to GEMINI_MODEL=gemini-2.5-flash. Failures are caught and converted to feature-specific fallback text.")
heading("Hosting and environment wiring", 2)
p("The production frontend default API URL is https://careledgerai.onrender.com; development defaults to http://127.0.0.1:8000. VITE_API_BASE_URL overrides either value at Vite build time. Vercel uses frontend/vercel.json to rewrite all client routes to index.html. The FastAPI CORS middleware accepts ALLOWED_ORIGINS as a comma-separated list with trailing slashes removed; without it, only localhost:5173 and 127.0.0.1:5173 are allowed. A Vercel deployment therefore needs its exact origin in the Render environment.")
code("Frontend: VITE_API_BASE_URL, VITE_SUPABASE_URL, VITE_SUPABASE_ANON_KEY\nBackend: GEMINI_API_KEY, GEMINI_MODEL (optional), ALLOWED_ORIGINS")

heading("3. Project Structure")
code("""CareledgerAI/
├─ frontend/
│  ├─ src/
│  │  ├─ App.jsx                 Active route map and landing page
│  │  ├─ main.jsx                React/Vite entry point
│  │  ├─ pages/                  Feature and authentication screens
│  │  ├─ components/             Shared UI, including ProtectedRoute
│  │  ├─ layouts/                Signed-in dashboard shell
│  │  ├─ store/useCareContext.js Zustand pipeline/session state
│  │  ├─ api/                    AI chat and Supabase chat-history clients
│  │  ├─ lib/api.js              Backend base-URL selection
│  │  ├─ lib/supabase.js         Supabase client and redirect helper
│  │  ├─ context/AuthContext.jsx Auth/session/profile-image state
│  │  └─ assets/                 Static illustrations and imagery
│  ├─ app/                       Older Expo/React Native scaffold; not used by Vite
│  ├─ package.json               Active web dependencies and Vite scripts
│  ├─ vite.config.js             React plugin configuration
│  ├─ tailwind.config.js         Tailwind theme/content configuration
│  └─ vercel.json                SPA fallback rewrite
├─ backend/
│  ├─ main.py                    FastAPI app, CORS, router registration
│  ├─ routers/                   HTTP schemas and route handlers
│  ├─ services/                  ML, rules, emergency, hospital, formatting logic
│  ├─ data/hospitals_pakistan.json  Curated demonstration catalog
│  ├─ ml/
│  │  ├─ train_model.py          Random-forest training/evaluation script
│  │  ├─ data/Training.csv       4,920-row training set
│  │  ├─ data/Testing.csv        42-row test set
│  │  ├─ disease_model.pkl       Serialized classifier
│  │  ├─ symptoms_list.pkl       Ordered 132-feature schema
│  │  ├─ diseases_list.pkl       41 disease labels
│  │  └─ knowledge_base/         Disease symptoms/specialist/severity rules
│  ├─ tests/                     Cost, affordability, and appointment tests
│  └─ requirements.txt           Python runtime dependencies
├─ supabase/migrations/          Chat tables, indexes, and RLS policies
├─ .env.example                  Deployment variable template
└─ README.md                     Short, partly outdated project overview""")

heading("4. Feature-by-Feature Logic")
heading("Health Assessment", 2)
p("Frontend journey: the user answers nine questions covering intensity, breathing, fever, hydration, dizziness, duration, trend, daily activity, and skin changes. Each choice has 0–5 points and some choices have an urgent flag. The browser sums points, divides by the maximum possible points, and rounds to a 0–100 score. Urgent flags or score ≥80 yield urgent; ≥60 high; ≥35 moderate; otherwise mild. The result screen shows the gauge, recommendation, answers, and emergency/care actions.")
p("API: none. Although GET /health-assessment exists, this page never calls it; that endpoint returns a fixed example response and is not part of the current UI flow.")
code('GET /health-assessment → {"risk_level":"moderate","likely_conditions":["seasonal flu","allergic rhinitis"],"recommended_action":"Consult a general physician within 24 hours."}')
p("Data classification: entirely deterministic frontend logic. It is neither ML nor Gemini. On a care-navigation action it writes latestCareContext with disease='Quick check-in summary', specialist='General Physician', the derived risk/score, human-readable answer strings, and source='Health Assessment'. A notable implementation mismatch is that Navigator checks for riskLevel='critical', while Assessment emits 'urgent'; therefore its automatic critical redirect condition does not fire, although the Assessment result itself offers Emergency Mode.")

heading("Disease Prediction", 2)
p("Frontend journey: the screen first downloads and alphabetizes the actual symptom feature list. The user searches/selects at least two symptoms, requests a prediction, reviews the leading candidate plus two alternatives, confidence, specialist, severity, evidence, optional follow-up questions, and any urgent warning. Selecting 'Find specialists' continues to Navigator. The selected symptoms/result are restored for the browser session.")
code('GET /health-assessment/symptoms-list\nResponse: {"symptoms":["itching","skin_rash", ...]}')
code('POST /health-assessment/predict-disease\nRequest: {"symptoms":["high_fever","cough"],"follow_up_answers":{"duration":"1 to 3 days"}}\nResponse: {"selected_symptoms":[...],"predictions":[{"disease":"...","ml_probability":0.0,"rule_score":0.0,"final_score":0.0,"probability":0.0,"confidence":0,"matched_symptoms":[],"missing_symptoms":[],"recommended_specialist":"...","specialist":"...","severity":"...","reason":"..."}],"best_prediction":{...},"follow_up_questions":[...],"confidence_threshold_met":false,"follow_up_answers":{...},"emergency":false,"urgent_warning":null}')
p("Backend logic: Pydantic rejects fewer than two distinct nonblank symptoms. The runtime loads disease_model.pkl, symptoms_list.pkl, and diseases_list.pkl once. Input labels are lowercased, spaces/hyphens become underscores, and matching positions in a 132-element zero vector become 1. A pandas DataFrame preserves training-column order. predict_proba is ranked and the top five model classes proceed to the rule engine. For each disease, the rule score is symptom coverage plus a 0.1 match bonus capped at 1. The final score is ml_probability + (1 - ml_probability) × 0.3 × rule_score. Results are re-ranked and the first five returned; the UI emphasizes the top three. If the best score is below 0.75, five or six fixed follow-up questions are returned. Follow-up free text is appended to the ML input, but because it rarely equals a feature name it is normally unmatched; emergency detection uses only the original selected symptom IDs.")
p("Data classification: the classifier output is real computed ML; rules/specialists/severity are curated knowledge-base data; follow-up questions and emergency rules are deterministic; Gemini is not used. Confidence is a hybrid score, not a calibrated clinical probability. The frontend saves the leading disease/specialist/confidence/risk/symptoms to latestCareContext.")

heading("Smart Care Navigator", 2)
p("Frontend journey: the page reads route state or latestCareContext. Without either it asks the user to complete Disease Prediction. With context, it calls the personalized endpoint, shows the condition/specialist, three top hospital/doctor cards, provider comparison, a cost preview, visit preparation, and booking links. Hospital selections are remembered for Cost, and hospital+doctor selections for Appointments.")
code('POST /smart-care-navigator/personalized\nRequest: {"disease":"Migraine","specialist":"Neurologist","risk_level":"medium"}\nResponse: {"featured_hospital":{...},"featured_doctor":{...},"nearby_hospitals":[...],"doctors":[...],"suggested_specialty":"Neurologist","cost_preview":{"min":0,"max":0},"what_to_expect":[...],"questions_to_ask":[...]}')
code('GET /smart-care-navigator → {"suggested_specialty":"General Medicine","nearby_hospitals":[...],"doctors":[...]}')
p("Backend logic: specialty names are normalized through aliases. Doctor match score is 2 for exact normalized specialty, 1 for overlapping specialty words, or 0. Hospitals sort by best match, count of matches, rating, wait, then distance. Doctors sort by match, rating, then experience. The featured provider is the first result. A deterministic Consultation and labs cost estimate supplies cost_preview. Gemini attempts JSON visit-preparation content; malformed, short, missing-key, or unavailable responses fall back to fixed guidance.")
p("Data classification: ranking and cost preview are computed from curated JSON. Institution names are used for realism, but doctor profiles, fees, ratings, availability, distance, and waiting times are fabricated/not live. Only visit guidance may be real Gemini output.")

heading("Cost Intelligence / Finance Module", 2)
p("Frontend journey: disease and hospital are prefilled from latestCareContext and selectedHospital. The user chooses one of eight treatment types, optionally chooses a catalog hospital, and requests an estimate. The maximum estimate is copied into the affordability form. Income, savings, and insurance then produce a score, label, out-of-pocket amount, narrative, and three suggestions.")
code('GET /medical-cost-intelligence/hospitals → {"hospitals":[{"id":"...","name":"...","city":"...","consultation_fee":2500}]}')
code('POST /medical-cost-intelligence/estimate\nRequest: {"predicted_disease":"Migraine","treatment_type":"Consultation and labs","selected_hospital":"hospital-id-or-null"}\nResponse: {"predicted_disease":"Migraine","treatment_type":"Consultation and labs","selected_hospital":"Hospital Name","specialist":"Neurologist","matched_doctor":"Doctor Name","cost_breakdown":{"consultation":{"min":0,"max":0},"labs":{"min":0,"max":0},"medication":{"min":0,"max":0},"procedure":{"min":0,"max":0},"total_range":{"min":0,"max":0}}}')
code('POST /medical-cost-intelligence/affordability-analysis\nRequest: {"total_cost_estimate":50000,"monthly_income":100000,"existing_savings":10000,"insurance_coverage_percent":20,"cost_breakdown":{},"treatment_type":"...","predicted_disease":"...","selected_hospital":"..."}\nResponse: {"affordability_score":58,"score_label":"Manageable","effective_out_of_pocket_cost":30000.0,"ai_summary":"...","ai_suggestions":["...","...","..."]}')
p("Backend logic: keyword rules choose a specialist. Treatment bands supply labs, medication, and procedure min/max PKR values. A selected hospital scales those bands by its consultation fee/2500, clamped to 0.65–1.45; its best matching doctor fee becomes the fixed consultation cost. Totals are sums of component bounds. Affordability computes insured cost, subtracts savings (floor zero), starts at 100, subtracts up to 60 for out-of-pocket/monthly-income pressure and up to 20 based on savings coverage, rounds/clamps, then labels ≥80 Easily affordable, ≥55 Manageable, otherwise Significant strain.")
p("Data classification: estimates and score are real deterministic calculations over illustrative bands/catalog data, not live quotes, claims, or bills. Gemini only writes the narrative/suggestions and receives no authority to change the numeric result. Fallback suggestions are deterministic.")

heading("Appointments", 2)
p("Frontend journey: the catalog loads, optionally ranked for a Navigator-recommended specialty. The user can sort hospitals, filter doctors, select a doctor, pick a future date supported by available_days, select a catalog time slot and In-person or Video call, enter a reason, and submit. Existing process-local bookings load on page entry and successful bookings appear immediately.")
code('GET /appointments/catalog?recommended_specialist=Neurologist\nResponse: {"hospitals":[...],"doctors":[{...,"is_recommended":true,"specialty_match_score":2}],"recommended_specialist":"Neurologist"}')
code('GET /appointments/my-appointments\nResponse: [{"appointment_id":"apt-...","hospital_id":"...","doctor_id":"...","hospital_name":"...","doctor_name":"...","specialization":"...","consultation_fee":2500.0,"appointment_date":"YYYY-MM-DD","time_slot":"10:00 AM","appointment_type":"In-person","reason_for_visit":"...","status":"Confirmed","confirmation_message":"..."}]')
code('POST /appointments/book\nRequest: {"hospital_id":"...","doctor_id":"...","appointment_date":"YYYY-MM-DD","time_slot":"...","appointment_type":"In-person","reason_for_visit":"..."}\nResponse: one AppointmentRecord with the fields shown above')
p("Backend logic: hospital/doctor IDs must exist together; the date must parse as YYYY-MM-DD and not be in the past; the weekday and slot must occur in the doctor profile; type must be exactly In-person or Video call; and duplicate doctor/date/slot combinations are rejected. An apt- plus eight-hex-character ID is generated. Gemini may phrase the confirmation; the deterministic fallback includes the actual booking data.")
p("Data classification: catalog and availability are sample data, validation/conflict detection are real, and confirmation prose may be Gemini. _session_bookings is a module-level list shared by all callers in one backend process, with no user association, persistence, cancellation, or update API.")

heading("AI Chat", 2)
p("Frontend journey: a signed-in user starts or reopens a chat, sends a message, sees an optimistic user bubble, receives the assistant reply, and may delete a whole session. The client tries to persist both sides in Supabase but still sends to Gemini if history backup fails. Any response with emergency_detected displays an urgent banner linking to Emergency Mode.")
code('POST /ai-chat/message\nRequest: {"message":"I have a headache","conversation_history":[{"role":"user","content":"..."},{"role":"assistant","content":"..."}],"user_name":"First name or null"}\nResponse: {"reply":"...","emergency_detected":false}')
p("Backend logic: message must be nonblank. System history entries are discarded; assistant/model roles become model and everything else becomes user. Regexes detect chest pain, breathing difficulty, severe bleeding, unconsciousness, and stroke phrases in the new message. Gemini receives a personalized, safety-constrained instruction, maximum 600 output tokens, and temperature 0.4. Empty/error responses use a calm fallback, prepending emergency advice when required.")
p("Data classification: the main reply is real Gemini output when configured and reachable; the emergency flag is deterministic keyword matching, not Gemini classification. Supabase chat_sessions/chat_messages are the only durable application content in the current implementation. RLS lets users manage their own sessions and related messages.")

heading("Emergency Mode", 2)
p("Frontend journey: the user opens a red emergency screen, can tap tel:1122, and expand fixed first-aid cards for burns, cuts, fainting, breathing difficulty, and choking. The page repeats that it is not a substitute for professional care and identifies the number as Pakistan-specific.")
p("API and backend: none. Emergency Mode is static UI. Entry points include urgent Assessment results, Disease Prediction emergency metadata, Navigator warnings, and AI Chat emergency flags. Disease Prediction emergency rules cover loss of consciousness; paralysis/one-sided weakness/slurred speech; certain bleeding symptoms; severe headache with neurological signs; and chest pain plus breathlessness. AI Chat uses separate natural-language regexes.")
p("Data classification: fixed authored guidance only—no live ambulance dispatch, location lookup, hospital availability, clinical triage service, or Gemini call.")

heading("5. Data Flow Between Features")
p("useCareContext is a Zustand store persisted under careledger-latest-care-context in sessionStorage. It survives reloads in the same tab/session but is neither server-side nor shared across devices.")
code("""latestCareContext: { disease, specialist, riskLevel, confidence, symptoms, source }
  Set by Assessment when the user continues; set automatically by Disease Prediction.
  Read by Navigator and Cost Intelligence.

selectedHospital: { id, name }
  Set by Navigator on recommended/selected providers.
  Read by Cost Intelligence to prefill selected_hospital.

appointmentSelection: { hospital_id, doctor_id, recommended_specialist }
  Set by Navigator booking links.
  Read by Appointments to preselect and specialty-rank its catalog.

diseasePredictionSession: { selectedSymptoms, result }
  Updated by Disease Prediction as inputs/results change.
  Read by Disease Prediction on remount to restore its screen.""")
p("Clear actions exist for every field but are not a substitute for logout cleanup; the store configuration itself does not clear on sign-out. Route location.state duplicates some handoffs for immediate navigation, while the persisted store provides reload resilience.")

heading("6. Authentication Flow")
numbered([
    "Initialization: AuthProvider calls supabase.auth.getSession(), subscribes to onAuthStateChange, and exposes user/session/loading. The client persists sessions, refreshes tokens, and detects sessions in callback URLs.",
    "Sign-up: name/email/password are validated in the browser; signUp stores full_name in user metadata. If Supabase returns a session, the user goes to /dashboard; otherwise the UI asks for email confirmation.",
    "Sign-in: signInWithPassword runs after nonblank validation, then navigates to /dashboard. Common provider, confirmation, and invalid-credential errors receive friendlier messages.",
    "Google OAuth: signInWithOAuth uses provider=google, redirects to the current origin/dashboard, and requests offline access with a consent prompt. Supabase/provider dashboard configuration must authorize the deployed callback.",
    "Password reset: resetPasswordForEmail redirects to /reset-password. That page waits for PASSWORD_RECOVERY or a recovery session in the URL, validates a six-character matching password, calls updateUser, then signs out.",
    "Route protection: /dashboard, /assessment, /disease-prediction, /navigator, /emergency, /cost, /appointments, /chat, /roadmap, and /profile are inside ProtectedRoute. The landing/sign-in/sign-up/forgot pages are public/guest-only; /reset-password is public so emailed recovery links work."
])
p("Protection is frontend-only. Backend requests carry no Supabase access token, FastAPI has no authentication dependency, and every backend endpoint—including appointment listings/bookings and AI chat—is callable without verified identity. Supabase independently protects chat rows through RLS. Profile metadata uses Supabase Auth; profile images are base64-like browser localStorage values keyed by user ID.")

heading("7. The ML Model (Disease Prediction)")
heading("Dataset and provenance", 2)
p("The repository includes backend/ml/data/Training.csv (4,920 rows) and Testing.csv (42 rows), with 132 ordered symptom feature columns and a prognosis target spanning 41 labels. No README, comment, license, URL, or attribution file identifies the original dataset source. It would be inaccurate to claim a source such as Kaggle without adding provenance evidence to the repository.")
heading("Training process", 2)
numbered([
    "Read both CSVs with pandas and remove columns whose names begin Unnamed:.",
    "Require prognosis in both datasets; derive symptom columns from training and verify train/test feature parity.",
    "Reorder test data to prognosis followed by training symptom order; sort unique disease labels.",
    "Fit RandomForestClassifier(n_estimators=200, random_state=42) on the binary symptom matrix.",
    "Predict the held-out Testing.csv rows and print accuracy_score.",
    "Serialize disease_model.pkl, symptoms_list.pkl, and diseases_list.pkl with joblib."
])
p("Evaluation of the checked-in artifact against the checked-in Testing.csv produced 0.9762 accuracy (41 correct out of 42). This is a tiny, apparently one-example-per-class style test set; it is not clinical validation, calibration evidence, or proof of real-world generalization.")
heading("Runtime use", 2)
p("Artifacts are loaded through an LRU-cached function, so one worker process loads them once. A normalized symptom-to-column index creates an exact binary vector; unmatched values are tracked internally. predict_proba ranks five classes. The rule engine enriches each with specialist/severity/evidence and slightly raises—never lowers—the probability according to curated symptom coverage. The API's confidence percentage is the rounded hybrid final_score.")
heading("Known ML limitations", 2)
bullets([
    "No dataset provenance or license is recorded in the repository.",
    "Only binary presence is modeled; age, sex, history, duration, severity, vitals, prevalence, medication, and negative symptoms are absent.",
    "The test set has only 42 rows for 41 labels, making the headline accuracy fragile.",
    "Random-forest probabilities and the hybrid rule uplift are not clinically calibrated.",
    "Exact vocabulary matching means natural-language follow-up answers usually do not activate features.",
    "The checked-in pickle reports an estimator serialization-version mismatch in the current local environment (created with scikit-learn 1.9.0, loaded under 1.7.2 during verification); requirements are unpinned, so reproducibility and compatibility are not guaranteed.",
    "The model and knowledge base provide decision support only and have no documented clinical review, bias testing, monitoring, or regulatory validation."
])

heading("8. Known Limitations / Honest Gaps")
bullets([
    "Appointments are stored in a module-level list. They vanish on restart, diverge across multiple workers/instances, are visible collectively rather than per user, and have no cancellation/rescheduling endpoints.",
    "FastAPI performs no authentication or authorization and does not validate Supabase JWTs. Frontend route guards are usability controls, not API security.",
    "Hospital names are real for demonstration realism, but doctors, fees, ratings, distance, wait times, insurance, availability, and related details are curated/fabricated and not live verified data.",
    "Cost bands are illustrative PKR planning estimates, not live provider quotes, invoices, benefit checks, claims, or financial advice.",
    "Health Assessment is a browser-only scoring questionnaire. Its GET backend endpoint is fixed sample output and unused by the screen.",
    "Navigator's critical Assessment redirect checks for 'critical', but Assessment emits 'urgent'; the automatic redirect branch is therefore inconsistent.",
    "Disease follow-up answers are appended as raw strings to an exact symptom vocabulary, so they generally do not change the ML vector. Emergency detection ignores those follow-up answers.",
    "Gemini failures degrade gracefully but are mostly hidden from users; appointment and navigator fallbacks may look like AI content even when deterministic.",
    "AI Chat emergency detection is a short regex list and can miss paraphrases or produce false positives. It is not a triage model.",
    "Emergency Mode is static and Pakistan-specific. It does not detect location, call on the user's behalf, or verify real-time emergency/provider availability.",
    "Supabase currently stores auth and chat only. Health assessments, predictions, care context, selected providers, costs, and appointments are not durable clinical records.",
    "Zustand data is sessionStorage data and is not explicitly cleared on logout, so another login in the same browser session could inherit prior care context.",
    "The README describes an Expo frontend and mock backend routes, while the active product is the Vite web app with a mixture of computed, AI, and sample-backed routes.",
    "Dependencies are mostly unpinned in backend/requirements.txt; deployment reproducibility is weak, especially for serialized scikit-learn artifacts.",
    "There is no documented audit log, consent workflow, encryption design beyond platform defaults, clinical-data retention policy, rate limiting, abuse protection, observability, or production health-check configuration."
])

# Footer and metadata
doc.core_properties.title = "CareLedger AI — Technical Documentation"
doc.core_properties.subject = "Verified implementation architecture and feature logic"
doc.core_properties.author = "CareLedger AI Engineering"
for sec in doc.sections:
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("CareLedger AI — Technical Documentation")

doc.save(OUTPUT)
print(f"Created {OUTPUT} ({OUTPUT.stat().st_size} bytes)")
